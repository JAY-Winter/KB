# import sys
# import os
# # 현재 스크립트 파일의 절대 경로를 가져옵니다.
# current_directory = os.path.dirname(os.path.abspath(__file__))

# # 해당 경로를 sys.path에 추가합니다.
# if current_directory not in sys.path:
#     sys.path.append(current_directory)

from fastapi import FastAPI, File
from fastapi.responses import FileResponse

from fastapi import FastAPI, UploadFile, HTTPException, Form, Depends
from fastapi.templating import Jinja2Templates
from fastapi.staticfiles import StaticFiles

from starlette.requests import Request

from sklearn.feature_extraction.text import TfidfVectorizer
from sklearn.metrics.pairwise import cosine_similarity

from gensim.summarization import summarize

import openai
from io import BytesIO
from docx import Document

from transformers import pipeline

app = FastAPI()

app.mount("/static", StaticFiles(directory="static"), name="static")

templates = Jinja2Templates(directory="templates")

openai.api_key = API_KEY

import spacy
from io import BytesIO
from docx import Document
import pytextrank

# spaCy 모델 로딩
nlp = spacy.load("en_core_web_sm")

# pytextrank를 파이프라인에 추가
tr = pytextrank
nlp.add_pipe("textrank")



def read_docx(file):
    doc = Document(file)
    return ' '.join(paragraph.text for paragraph in doc.paragraphs)

@app.get("/")
async def read_form(request: Request):
    return templates.TemplateResponse("upload.html", {"request": request})

@app.post("/summary/")
async def summarize_docx(file: UploadFile = File(...)):
    if file.filename.endswith('.docx'):
        file_contents = await file.read()
        document = read_docx(BytesIO(file_contents))
        doc = nlp(document)
        print('doc')
        print(doc)
        # 문장 중요도 순서로 문장들을 정렬
        # sentences = [sent.text for sent in doc._.textrank.summary(limit_phrases=15, limit_sentences=5)]
        sentences = [sent.text for sent in doc._.textrank.summary()]
        summary = ' '.join(sentences)
        # 상위 n 문장을 선택하여 요약 생성 (여기서는 상위 5문장을 사용)
        print()
        print('summary')
        print(summary)

        # Gensim 요약
        # summarized_text = summarizdocumente(document, word_count=100)
        # print('summarized_text', summarized_text)
        # GPT-3 요약
        # gpt_response = summary_GPT(summarized_text)
        gpt_response = summary_GPT(summary)
            
        # 유사도 비교        
        # 업로드된 파일과 로컬에 저장된 파일들 사이의 유사도를 비교하여 정렬된 파일 리스트를 가져옵니다.
        similar_documents = compare_similarity(document)
        print(similar_documents)

        # 최종 결과 반환
        return {"gpt_summary": gpt_response, 'similar_documents': similar_documents}
    
    else:
        raise HTTPException(status_code=400, detail="Invalid file type")


@app.get('/file/')
def get_file(path):
    return FileResponse(path, filename=path)


def docx_to_text(file_path):
    doc = Document(file_path)
    return ' '.join(paragraph.text for paragraph in doc.paragraphs)


def compare_similarity(uploaded_file_content):
    # 로컬에 저장된 문서들을 로드합니다.
    file_paths = ['/Users/heyon/Desktop/KB/file/1.docx', '/Users/heyon/Desktop/KB/file/2.docx', '/Users/heyon/Desktop/KB/file/3.docx']
    documents = [docx_to_text(file_path) for file_path in file_paths]

    # 업로드된 파일의 내용도 리스트에 추가합니다.
    documents.append(uploaded_file_content)

    # TfidfVectorizer를 사용하여 문서들을 벡터화합니다.
    vectorizer = TfidfVectorizer()
    tfidf_matrix = vectorizer.fit_transform(documents)

    # 마지막 문서 (업로드된 문서)와 나머지 문서들의 유사도를 계산합니다.
    similarities = cosine_similarity(tfidf_matrix[-1], tfidf_matrix[:-1])

    # 유사도와 파일 경로를 결합하고 유사도 순으로 정렬합니다.
    sorted_files = sorted(zip(similarities[0], file_paths), key=lambda x: x[0], reverse=True)

    # 정렬된 파일 경로를 반환합니다.
    return [file for file in sorted_files]




# 한국어에 적합한 TF-IDF 추출

# from konlpy.tag import Okt
# from sklearn.feature_extraction.text import TfidfVectorizer
# import re

# Okt 형태소 분석기 인스턴스 생성
# okt = Okt()

# # 간단한 한국어 stop words 리스트
# # 이 리스트는 원하는대로 확장/수정할 수 있습니다.
# stop_words = ["의", "가", "이", "은", "들", "는", "좀", "잘", "걍", "과", "도", "를", "으로", "자", "에", "와", "한", "하다"]

# def extract_important_sentences(text, num_keywords=5, num_sentences=5):
#     vectorizer = TfidfVectorizer(max_df=0.85, stop_words=stop_words, max_features=num_keywords, tokenizer=okt.morphs)
#     tfidf_matrix = vectorizer.fit_transform([text])
#     feature_names = vectorizer.get_feature_names_out()
    
#     sentences = re.split(r'(?<=[.!?])\s+', text)
#     important_sentences = []
    
#     for sentence in sentences:
#         if any(word in sentence for word in feature_names):
#             important_sentences.append(sentence)
#             if len(important_sentences) >= num_sentences:
#                 break
    
#     return ' '.join(important_sentences)


# GPT-3 요약

def request_openai_summary(prompt_text):
    try:
        response = openai.Completion.create(
            model="text-davinci-003",
            prompt=prompt_text,
            temperature=0.7,
            max_tokens=2000,
            top_p=1,
            frequency_penalty=0.0,
            presence_penalty=0.0,
            stop=["\n"]
        )
        return response
    except openai.error.InvalidRequestError as e:
        return None


def summary_GPT(text):
    prompt_text = "한국어로 최소 1문장, 최대 50문장으로 다음 텍스트 요약하는데, '어떤 문서입니다' 식으로 요약해줘" + "----" + text
    response = request_openai_summary(prompt_text)
    
    if not response:
        # re_prompt_text = extract_important_sentences(text)
        re_prompt_text = text[:100]
        print('re_prompt')
        print(re_prompt_text)
        response = request_openai_summary(re_prompt_text)

    gpt_response = response.choices[0].text.strip() if response else '별도의 이유로 요약되지 못 했습니다.'
    return gpt_response