import os
import aiofiles
from pprint import pprint
from docx import Document
from fastapi import HTTPException
from sklearn.metrics.pairwise import cosine_similarity
from sklearn.feature_extraction.text import TfidfVectorizer


def get_all_txt_files_in_directory(directory_path):
    '''
    주어진 디렉토리 내의 모든 .txt 파일의 경로를 리스트로 반환합니다.
    '''
    return [os.path.join(directory_path, file) for file in os.listdir(directory_path) if file.endswith('.txt')]


def txt_to_text(file_path):
    '''
    주어진 .txt 파일의 내용을 문자열로 반환합니다.
    '''
    with open(file_path, 'r', encoding='utf-8') as file:
        return file.read()


def read_docx(file_path: str) -> str:
    '''
    주어진 경로의 .docx 파일 내용을 문자열로 반환합니다.
    '''
    doc = Document(file_path)
    return ' '.join(paragraph.text for paragraph in doc.paragraphs if paragraph.text.strip() != '')


def docx_to_text(file_path):
    doc = Document(file_path)
    return ' '.join(paragraph.text for paragraph in doc.paragraphs)


def read_file_from_path(path: str) -> str:
    '''
    주어진 경로에서 파일을 읽어 텍스트를 반환합니다.
    '''
    try:
        if path.endswith('.docx'):
            return read_docx(path)
        if path.endswith('.txt'):
            with open(path, mode='r', encoding='utf-8') as f:
                return f.read()
    except Exception as e:
        raise HTTPException(status_code=400, detail=str(e))


def get_all_files_in_directory(search_directory_path, file_type):
    '''
    주어진 디렉토리 내의 파일 타입별 모든 파일의 경로를 리스트로 반환합니다.
    '''
    if file_type == 'docx':
        return [os.path.join(search_directory_path, file) for file in os.listdir(search_directory_path) if file.endswith('.docx')]
    elif file_type == 'txt':
        return [os.path.join(search_directory_path, file) for file in os.listdir(search_directory_path) if file.endswith('.txt')]
    return [os.path.join(search_directory_path, file) for file in os.listdir(search_directory_path)]


async def compare_similarity(uploaded_file_name, uploaded_file_content, file_type, search_directory_path):
    '''
    업로드한 파일과 로컬 파일의 유사도를 비교합니다.
    '''
    from module.SINGLETONE import MainInstance as MainInstance
    from main import Redis_Instance

    # Redis 를 활용한 Cache 적용
    redis_client = Redis_Instance.redis_client
    hash_key = 'uploadAPI' + search_directory_path + '/' +uploaded_file_name

    cached_result = redis_client.get(hash_key)

    if cached_result:
        return eval(cached_result.decode('utf-8'))
    else:
        all_files = MainInstance.get_instance().get_all_files()
        documents = MainInstance.get_instance().get_documents()
        vectorizer = MainInstance.get_instance().get_vectorizer()
        tfidf_matrix = MainInstance.get_instance().get_tfidf_matrix()
        uploaded_file_vector = vectorizer.transform([uploaded_file_content])

        # 마지막 문서 (업로드된 문서)와 기존 문서들의 유사도를 계산합니다.
        similarities = cosine_similarity(uploaded_file_vector, tfidf_matrix)

        # 유사도와 파일 경로를 결합하고 유사도 순으로 정렬합니다.
        sorted_files = sorted(zip(similarities[0], all_files), key=lambda x: x[0], reverse=True)

        # 정렬된 파일 10개를 반환합니다.
        RETURN_CNT = 10
        results = [file for file in sorted_files[:RETURN_CNT]]
        
        # Redis 에 해당 파일과 연관된 파일 값을 저장합니다
        redis_client.set(hash_key, str(results), ex=3600)
        return results